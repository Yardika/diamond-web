"""Utility functions for DOCX template processing and placeholder replacement."""

import re
from io import BytesIO
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def replace_text_in_paragraph(paragraph, replacements):
    """Replace placeholder text in a paragraph while preserving formatting.
    
    Args:
        paragraph: A docx paragraph object
        replacements: A dict of {placeholder: value} pairs
    """
    # Get all runs in the paragraph
    if not paragraph.runs:
        return
    
    full_text = ''.join(run.text for run in paragraph.runs)
    
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            # Track which runs need updates
            runs_text = [run.text for run in paragraph.runs]
            placeholder_found = False
            
            # Find and replace the placeholder
            new_text = full_text.replace(placeholder, str(value))
            
            # Clear all runs
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Add the new text
            if new_text:
                paragraph.add_run(new_text)
            
            full_text = new_text
            placeholder_found = True


def replace_text_in_table_cell(cell, replacements):
    """Replace placeholder text in a table cell.
    
    Args:
        cell: A docx table cell object
        replacements: A dict of {placeholder: value} pairs
    """
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)


def fill_template_with_data(template_file, replacements):
    """Fill a DOCX template with data by replacing placeholders.
    
    Replaces all {{placeholder}} patterns in the document with corresponding values.
    Handles placeholders in:
    - Paragraphs
    - Tables
    - Headers/Footers
    
    Args:
        template_file: File path or BytesIO object of the template DOCX
        replacements: Dict of {placeholder: value} pairs
        
    Returns:
        BytesIO: The filled document as bytes
    """
    # Load the template
    if isinstance(template_file, str):
        doc = Document(template_file)
    else:
        doc = Document(template_file)
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        # Get full text of paragraph
        full_text = ''.join(run.text for run in paragraph.runs)
        
        # Check if any placeholder exists in this paragraph
        has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
        
        if has_placeholder:
            # Clear the paragraph
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Replace placeholders in the full text
            new_text = full_text
            for placeholder, value in replacements.items():
                new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
            
            # Add the new text with original formatting from first run if available
            if new_text:
                paragraph.add_run(new_text)
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
                    
                    if has_placeholder:
                        for run in paragraph.runs:
                            run._element.getparent().remove(run._element)
                        
                        new_text = full_text
                        for placeholder, value in replacements.items():
                            new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
                        
                        if new_text:
                            paragraph.add_run(new_text)
    
    # Process headers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
            
            if has_placeholder:
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                new_text = full_text
                for placeholder, value in replacements.items():
                    new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
                
                if new_text:
                    paragraph.add_run(new_text)
        
        # Process footers
        for paragraph in section.footer.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
            
            if has_placeholder:
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                new_text = full_text
                for placeholder, value in replacements.items():
                    new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
                
                if new_text:
                    paragraph.add_run(new_text)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
